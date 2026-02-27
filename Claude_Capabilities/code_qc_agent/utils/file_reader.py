"""File reading utilities for source-of-truth and code files."""

from pathlib import Path


def read_sot_file(path: str) -> str:
    """Read a source-of-truth file and return its content as text.

    Supported formats: .xlsx, .docx, .pdf
    """
    p = Path(path)
    ext = p.suffix.lower()

    if ext == ".xlsx":
        return _read_xlsx(p)
    elif ext == ".docx":
        return _read_docx(p)
    elif ext == ".pdf":
        return _read_pdf(p)
    else:
        raise ValueError(f"Unsupported source-of-truth format: {ext!r}. Use .xlsx, .docx, or .pdf")


def _read_xlsx(path: Path) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(path, data_only=True)
    parts: list[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"=== Sheet: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            row_text = "\t".join(str(v) if v is not None else "" for v in row)
            if row_text.strip():
                parts.append(row_text)

    return "\n".join(parts)


def _read_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        parts.append("[TABLE]")
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    return "\n".join(parts)


def _read_pdf(path: Path) -> str:
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            text = page.extract_text()
            if text:
                parts.append(f"=== Page {i} ===")
                parts.append(text)

    return "\n".join(parts)


def read_code_file(path: str) -> tuple[str, str]:
    """Read a code file and return (raw_code, language).

    Supported languages: .py → python, .R/.r → r, .sas → sas
    """
    p = Path(path)
    ext = p.suffix.lower()

    language_map = {
        ".py": "python",
        ".r": "r",
        ".sas": "sas",
    }

    language = language_map.get(ext)
    if language is None:
        raise ValueError(
            f"Unsupported code file format: {ext!r}. Use .py, .R, or .sas"
        )

    raw_code = p.read_text(encoding="utf-8", errors="replace")
    return raw_code, language
