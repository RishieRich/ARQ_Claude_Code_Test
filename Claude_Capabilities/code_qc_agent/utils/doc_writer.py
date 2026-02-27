"""Word document report writer for the Code QC Agent."""

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils.models import LogicQCResult, StructureQCResult, QCFinding


# Hex fill colours for status cells
_STATUS_BG = {
    "PASS": "00B050",     # Green
    "FAIL": "C00000",     # Red
    "WARNING": "FF8C00",  # Orange
}

# White text on FAIL cells for contrast
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_BLACK = RGBColor(0x00, 0x00, 0x00)

_HEADER_BG = "2F5496"   # Dark blue for table headers


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, color_hex: str) -> None:
    """Set a table cell background colour via XML shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove any existing shd element to avoid duplicates
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    tcPr.append(shd)


def _set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None) -> None:
    """Clear a cell and write text with optional formatting."""
    cell.text = text
    if not cell.paragraphs[0].runs:
        return
    run = cell.paragraphs[0].runs[0]
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _add_header_row(table, headers: list[str]) -> None:
    """Style the first row of a table as a blue header."""
    hdr_row = table.rows[0]
    for i, label in enumerate(headers):
        cell = hdr_row.cells[i]
        _set_cell_bg(cell, _HEADER_BG)
        _set_cell_text(cell, label, bold=True, color=_WHITE)


# ---------------------------------------------------------------------------
# Findings table
# ---------------------------------------------------------------------------

def _add_findings_table(doc: Document, findings: list[QCFinding]) -> None:
    if not findings:
        doc.add_paragraph("No findings recorded.")
        return

    headers = ["Title", "Status", "Detail", "Recommendation", "Priority"]
    table = doc.add_table(rows=1 + len(findings), cols=len(headers))
    table.style = "Table Grid"
    _add_header_row(table, headers)

    for row_idx, finding in enumerate(findings, start=1):
        values = [
            finding.title,
            finding.status,
            finding.detail,
            finding.recommendation,
            finding.priority,
        ]
        for col_idx, val in enumerate(values):
            cell = table.cell(row_idx, col_idx)
            cell.text = val

            # Colour the Status cell
            if col_idx == 1:
                bg = _STATUS_BG.get(finding.status, "FFFFFF")
                _set_cell_bg(cell, bg)
                if finding.status == "FAIL" and cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.color.rgb = _WHITE


# ---------------------------------------------------------------------------
# Executive summary table
# ---------------------------------------------------------------------------

def _add_summary_table(
    doc: Document,
    logic_result: LogicQCResult,
    structure_result: StructureQCResult,
) -> None:
    total_pass = logic_result.pass_count + structure_result.pass_count
    total_fail = logic_result.fail_count + structure_result.fail_count
    total_warn = logic_result.warning_count + structure_result.warning_count

    rows_data = [
        ("Logic QC",     logic_result.pass_count,     logic_result.fail_count,     logic_result.warning_count),
        ("Structure QC", structure_result.pass_count, structure_result.fail_count, structure_result.warning_count),
        ("Total",        total_pass,                  total_fail,                  total_warn),
    ]

    table = doc.add_table(rows=1 + len(rows_data), cols=4)
    table.style = "Table Grid"
    _add_header_row(table, ["Category", "PASS", "FAIL", "WARNING"])

    for row_idx, (label, passes, fails, warns) in enumerate(rows_data, start=1):
        table.cell(row_idx, 0).text = label

        pass_cell = table.cell(row_idx, 1)
        pass_cell.text = str(passes)
        _set_cell_bg(pass_cell, "00B050")

        fail_cell = table.cell(row_idx, 2)
        fail_cell.text = str(fails)
        if fails > 0:
            _set_cell_bg(fail_cell, "C00000")
            if fail_cell.paragraphs[0].runs:
                fail_cell.paragraphs[0].runs[0].font.color.rgb = _WHITE

        warn_cell = table.cell(row_idx, 3)
        warn_cell.text = str(warns)
        if warns > 0:
            _set_cell_bg(warn_cell, "FF8C00")


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def write_report(
    sot_path: str,
    code_path: str,
    logic_result: LogicQCResult,
    structure_result: StructureQCResult,
    output_path: str,
) -> None:
    """Build and save the QC Word report."""
    doc = Document()

    # --- Title ---
    title = doc.add_heading("Code QC Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Metadata ---
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Source of Truth: {Path(sot_path).name}")
    doc.add_paragraph(f"Code File:       {Path(code_path).name}")

    # --- Executive Summary ---
    doc.add_heading("Executive Summary", level=1)
    _add_summary_table(doc, logic_result, structure_result)
    doc.add_paragraph("")

    # --- Logic QC Findings ---
    doc.add_heading("Logic QC Findings", level=1)
    _add_findings_table(doc, logic_result.findings)
    doc.add_paragraph("")

    # --- Structure QC Findings ---
    doc.add_heading("Structure QC Findings", level=1)
    _add_findings_table(doc, structure_result.findings)

    # --- Save ---
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
